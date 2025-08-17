from datetime import date, datetime
from re import S
import os
from selenium import webdriver
from PIL import Image
from os import listdir
from os.path import isfile, join
import os
from docx import Document
from docx.shared import Inches
from robot.libraries.BuiltIn import BuiltIn
from docx2pdf import convert
import re
import unicodedata
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
import time

########################## GERANDO WORD DAS EVIDÊNCIAS DA AUTOMAÇÃO
def gerar_evidencia_com_passos_extraidos():
    nome_cenario = BuiltIn().get_variable_value("${TEST NAME}")
    nome_arquivo = BuiltIn().get_variable_value("${SUITE SOURCE}")

    nome_base = sanitizar_nome_arquivo(f"{nome_cenario}")
    passos = extrair_passos_do_cenario(nome_arquivo, nome_cenario)

    pasta_logs = '.logs\\'
    pasta_saida = extrair_e_criar_diretorio_de_contexto(nome_cenario)

    imagens = [
        os.path.join(pasta_logs, f)
        for f in os.listdir(pasta_logs)
        if f.endswith('.png') or f.endswith('.jpg')
    ]
    imagens.sort(key=lambda x: os.path.getmtime(x))

    doc = Document()
    doc.add_heading('✅  Evidências de Testes Automatizados - Magento', level=1)
    doc.add_paragraph(f'{nome_cenario}')
    doc.add_paragraph(f'📅  Data/Hora: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
    doc.add_heading('📝  Passos do Teste', level=2)
    doc.add_paragraph(passos if passos else "Passos não encontrados.")

    if imagens:
        doc.add_page_break()
        doc.add_heading('📜 Evidências Listadas abaixo: ', level=2)
        for img in imagens:
            doc.add_picture(img, width=Inches(6.5), height=Inches(4.8))
            doc.add_page_break()

    caminho_arquivo_docx = os.path.join(pasta_saida, f"{nome_base}.docx")
    doc.save(caminho_arquivo_docx)
    print(f"✅ Evidência gerada: {caminho_arquivo_docx}")

    # Converte o DOCX em PDF
    converter_docx_para_pdf(caminho_arquivo_docx)
    
    for img in imagens:
        os.remove(img)

def extrair_passos_do_cenario(caminho_arquivo_robot, nome_cenario):
    with open(caminho_arquivo_robot, 'r', encoding='utf-8') as f:
        linhas = f.readlines()

    capturando = False
    passos = []
    for linha in linhas:
        linha_strip = linha.strip()
        if linha_strip.startswith(nome_cenario):
            capturando = True
            continue
        if capturando:
            if linha_strip.startswith("***") or linha_strip.lower().startswith("cenario"):
                break
            if linha_strip:
                passos.append(linha_strip)
    return "\n".join(passos)

def sanitizar_nome_arquivo(nome, max_length=150):
    nome = unicodedata.normalize('NFKD', nome).encode('ASCII', 'ignore').decode('ASCII')
    nome = re.sub(r"[<>:\"/\\|?*\-\[\]']", '', nome)
    nome = nome.replace(" ", "_")
    return nome[:max_length] if len(nome) > max_length else nome

def extrair_e_criar_diretorio_de_contexto(nome_cenario, base_path='resources\\Arquivos\\Evidencias'):
    # Extrai o texto entre colchetes []
    match = re.search(r'\[(.*?)\]', nome_cenario)
    if match:
        nome_pasta = sanitizar_nome_arquivo(match.group(1))
        caminho_completo = os.path.join(base_path, nome_pasta)
        os.makedirs(caminho_completo, exist_ok=True)
        return caminho_completo
    else:
        # Caso não tenha colchetes, salva direto na pasta base
        os.makedirs(base_path, exist_ok=True)
        return base_path

##### USAR COM O WORD INSTALADO
def converter_docx_para_pdf(caminho_arquivo_docx):
    if not os.path.exists(caminho_arquivo_docx):
        print(f"Arquivo DOCX não encontrado: {caminho_arquivo_docx}")
        return None

    if not caminho_arquivo_docx.lower().endswith(".docx"):
        print("O arquivo fornecido não é um .docx")
        return None

    caminho_pdf = caminho_arquivo_docx.replace(".docx", ".pdf")

    try:
        convert(caminho_arquivo_docx, caminho_pdf)
        print(f"✅ Arquivo PDF gerado em: {caminho_pdf}")
        return caminho_pdf
    except Exception as e:
        print(f"❌ Erro ao converter para PDF: {e}")
        return None

##### USAR COM O LIBREOFFICE INSTALADO
# def converter_docx_para_pdf(caminho_docx):
#     if not os.path.exists(caminho_docx):
#         print(f"❌ Arquivo não encontrado: {caminho_docx}")
#         return None

#     pasta_saida = os.path.dirname(os.path.abspath(caminho_docx))

#     # Caminho completo para o soffice.exe
#     soffice_path = r"C:\\Program Files\\LibreOffice\\program\\soffice.exe"
    
#     try:
#         # Comando para chamar o LibreOffice em modo headless (sem interface)
#         subprocess.run([
#             #"soffice",
#             soffice_path,
#             "--headless",
#             "--convert-to", "pdf",
#             "--outdir", pasta_saida,
#             caminho_docx
#         ], check=True)

#         caminho_pdf = os.path.splitext(caminho_docx)[0] + ".pdf"
#         print(f"✅ PDF gerado em: {caminho_pdf}")
#         return caminho_pdf
#     except subprocess.CalledProcessError as e:
#         print(f"❌ Erro ao converter para PDF: {e}")
#         return None


########################## GERANDO PRINT DOS LOGS DO ROBOT APÓS CADA EXECUÇÃO DA SUITE
def tirar_print_estatisticas_do_log(caminho_log_html, caminho_destino_png):
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    #chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--window-size=1280,1030")
    chrome_options.add_argument('--hide-scrollbars')

    driver = webdriver.Chrome(options=chrome_options)
    caminho_absoluto = 'file://' + os.path.abspath(caminho_log_html)
    driver.get(caminho_absoluto)

    time.sleep(3)  # aguarda o carregamento completo

    # Ajusta o tamanho da janela para caber todo o conteúdo
    largura = driver.execute_script("return document.body.scrollWidth")
    altura = driver.execute_script("return document.body.scrollHeight")
    driver.set_window_size(largura, altura)

    # Tira screenshot da tela inteira
    driver.save_screenshot(caminho_destino_png)
    driver.quit()
